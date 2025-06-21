import docx

from pathlib import Path

import logging
logging.basicConfig(level=logging.DEBUG, format="%(asctime)s - %(levelname)s - %(message)s")
def main():
    inpath = Path("target/infile.txt")
    wordpath = Path("target/invfile.docx")

    doc = docx.Document(f"{wordpath}")

    with open(inpath) as f:
        lines = f.readlines()
        logging.debug(lines)
        for line in lines:
            line = line.strip("\n")
            doc.add_paragraph("I would like to cordually invite you", style="Caption")
            doc.add_paragraph(f"{line}", style="Heading")
            doc.add_paragraph("to a meeting at 111 Spice Avenue, Arrakis", style="Caption")
            doc.add_paragraph("April 1st", style="Body Text")
            doc.add_paragraph("10pm", style="Caption")
            doc.add_page_break()
            print(line)
        doc.save("target/outfile.docx")




if __name__ == "__main__":
    main()
